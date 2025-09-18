from __future__ import annotations

import io
import re
from typing import Dict, List

from flask import (
    Flask,
    flash,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from PyPDF2 import PdfReader
from docx import Document


app = Flask(__name__)
app.secret_key = "change-me"


DEFAULT_CATEGORIES = ["Title", "Ingredients", "Instructions", "Notes"]


def split_into_sentences(block: str) -> List[str]:
    """Split a block of text into smaller sentence-like chunks."""
    bullet_pattern = re.compile(r"^([-*\u2022]\s+|\d+[).]\s+)")
    sentence_boundary = re.compile(r"(?<=[.!?])\s+(?=[A-Z0-9])")

    chunks: List[str] = []
    buffer: List[str] = []

    def flush_buffer() -> None:
        if not buffer:
            return
        text = " ".join(buffer)
        for piece in sentence_boundary.split(text):
            piece = piece.strip()
            if piece:
                chunks.append(piece)
        buffer.clear()

    for raw_line in block.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if bullet_pattern.match(line):
            flush_buffer()
            chunks.append(line)
        else:
            buffer.append(line)

    flush_buffer()
    return chunks


def extract_text_segments(file_stream: io.BytesIO) -> List[str]:
    """Extract cleaned text segments from the uploaded PDF."""
    reader = PdfReader(file_stream)
    segments: List[str] = []

    for page in reader.pages:
        raw_text = page.extract_text() or ""
        if not raw_text:
            continue

        normalized = raw_text.replace("\r", "\n")
        lines = [line.strip() for line in normalized.split("\n")]
        cleaned_lines = []

        for line in lines:
            if not line:
                cleaned_lines.append("")
                continue

            if re.match(r"^(page|p)\s*\d+\b", line, re.IGNORECASE):
                continue
            if re.match(r"^\d+$", line):
                continue

            cleaned_lines.append(line)

        joined = "\n".join(cleaned_lines)
        for block in re.split(r"\n{2,}", joined):
            block = block.strip()
            if block:
                segments.extend(split_into_sentences(block))

    return segments


@app.route("/", methods=["GET"])
def index():
    return render_template("upload.html")


@app.route("/extract", methods=["POST"])
def extract():
    uploaded = request.files.get("pdf")
    if not uploaded or uploaded.filename == "":
        flash("Please choose a PDF file to upload.")
        return redirect(url_for("index"))

    if not uploaded.filename.lower().endswith(".pdf"):
        flash("Only PDF files are supported.")
        return redirect(url_for("index"))

    file_bytes = io.BytesIO(uploaded.read())
    file_bytes.seek(0)

    try:
        segments = extract_text_segments(file_bytes)
    except Exception:  # pragma: no cover - PyPDF2 errors bubble up rarely
        flash("We couldn't read that PDF. Please try a different file.")
        return redirect(url_for("index"))

    if not segments:
        flash("No readable text was found in the PDF.")
        return redirect(url_for("index"))

    session["segments"] = segments
    session["categories"] = DEFAULT_CATEGORIES
    return render_template(
        "review.html",
        segments=segments,
        categories=DEFAULT_CATEGORIES,
    )


@app.route("/download", methods=["POST"])
def download():
    stored_segments = session.get("segments")
    if not stored_segments:
        flash("Please upload a PDF before requesting a download.")
        return redirect(url_for("index"))

    assignments: Dict[int, str] = {}
    positions: Dict[int, int] = {}
    for key, value in request.form.items():
        if key.startswith("assignment_"):
            try:
                index = int(key.split("_", 1)[1])
            except ValueError:
                continue
            assignments[index] = value
        elif key.startswith("position_"):
            try:
                index = int(key.split("_", 1)[1])
                positions[index] = int(value)
            except ValueError:
                continue

    kept_assignments = {
        idx: category
        for idx, category in assignments.items()
        if category not in {"discard", "unassigned"}
    }

    if not kept_assignments:
        flash("Assign at least one sentence to a category before downloading.")
        return render_template(
            "review.html",
            segments=stored_segments,
            categories=session.get("categories", DEFAULT_CATEGORIES),
        )

    sorted_indices = sorted(
        kept_assignments.keys(), key=lambda idx: positions.get(idx, idx)
    )

    ordered_categories: List[str] = []
    for idx in sorted_indices:
        category = kept_assignments[idx]
        if category not in ordered_categories:
            ordered_categories.append(category)

    document = Document()
    document.add_heading("Recipe", level=1)

    for category in ordered_categories:
        document.add_heading(category, level=2)
        for idx in sorted_indices:
            if kept_assignments[idx] == category:
                document.add_paragraph(stored_segments[idx])

    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return send_file(
        output,
        as_attachment=True,
        download_name="recipe.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


if __name__ == "__main__":
    app.run(debug=True)
